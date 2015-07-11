from html.parser import HTMLParser
from docx import Document
import docx
from docx.shared import Inches
from tools.Simple_WebCatcher import HTMLClient
import os
import re
class MyHTMLParser(HTMLParser):
    def __init__(self):
        HTMLParser.__init__(self)
        self.links = []
        self.attrs = []
 
    def handle_starttag(self, tag, attrs):
        print("Encountered the beginning of a %s tag" % tag)
#        if tag == "div":
#            if len(attrs) == 0: pass
#            else:
#                #print(attrs)
#                self.attrs = attrs
#                for (variable, value)  in attrs:
#                    if variable == "class":
#                        if value == "ImgSearch":
#                            print('find')
#                #return attrs
#    def handle_data(self, data):
#        print(data)
#        attrs = self.attrs
#        for (variable, value)  in attrs:
#            if variable == "class":
#                if value == "ImgSearch":
#                    print("data"+data)
    def handle_endtag(self, tag):
        print("Encountered the beginning of a %s tag" % tag)
#html = MyHTMLParser()
#html.feed('<p>abcd</br></p>')
#html = 'c:/html/kkk/start'
#print(html.split(r'/')[-1])

errorimg = '14c642b1f9ffac3fc406107e.png'
class XQHTMLParser(HTMLParser):
    def __init__(self, docfile):
        HTMLParser.__init__(self)
        self.docfile = docfile
        self.doc = Document(docfile)
        self.myclient = HTMLClient()
        self.text = ''
        self.title = False
        self.isdescription = False
        self.picList=[]
    def handle_starttag(self, tag, attrs):
        #print "Encountered the beginning of a %s tag" % tag
        self.title = False
        self.isdescription = False
        if re.match(r'h(\d)', tag):
            self.title = True 
        if tag == "img":
            if len(attrs) == 0: pass
            else:
                for (variable, value)  in attrs:
                    if variable == "src":
                        picdata = self.myclient.GetPic(value.split('!')[0])
                        if picdata == None:
                            pass
                        else:
                            pictmp = value.split('/')[-1].split('!')[0]
                            picfix = value.split('/')[-1].split('!')[-1]
                            with open(pictmp, 'wb') as pic:
                                pic.write(bytes(picdata))
                                pic.close()
                            #if os.path.getsize(pictmp) < 90000:
                            try:
                                if picfix[0:1] == 'c':
                                    self.doc.add_picture(pictmp, width=Inches(4.5))
                                else:
                                    self.doc.add_picture(pictmp)#, width=Inches(2.25))
                            except docx.image.exceptions.UnexpectedEndOfFileError as e:
                                print(e)
                            self.picList.append(pictmp)
        if tag == 'script':
            self.isdescription = True
    def handle_data(self, data):
        if self.title == True:
            if self.text != '':
                self.doc.add_paragraph(self.text)
            self.text = ''
            self.doc.add_heading(data, level=2)
        if self.isdescription == False:
            self.text += data
    def handle_endtag(self, tag):
        #if tag == 'br' or tag == 'p' or tag == 'div':
        if self.text != '':
            self.doc.add_paragraph(self.text)
            self.text = ''
    def complete(self, html):
        self.feed(html)
        self.doc.save(self.docfile)
        for item in self.picList:
            if os.path.exists(item):
                os.remove(item)

#xq_parser = XQHTMLParser("tmp.doc")
#xq_parser.feed(
#json = '<img src="http://xqimg.imedao.com/14b069587ccc53fc199a1ca2.jpg"  />'
#json = '<img src="http://xqimg.imedao.com/14b069587ccc53fc199a1ca2.jpg!custom.jpg" >'
#json = '<img src="http://xqimg.imedao.com/14b0696fb404e53fd90be7d0.jpg!custom.jpg" >'
#json = '<img src="http://xavatar.imedao.com/community/201411/1418865775807-1418865776034.jpeg!50x50.png">'
#xq_parser.complete(json)
#doc.save('tmp.doc')
